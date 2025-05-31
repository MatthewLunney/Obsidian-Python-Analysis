import tkinter as tk
from tkinter import ttk, messagebox
import pandas as pd
import datetime
import importlib.util
import sys
import os
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from collections import defaultdict

def select_sector_and_dates():
    root = tk.Tk()
    root.title("Select Sector and Date Range")
    root.geometry("400x250")
    root.resizable(False, False)

    xl = pd.ExcelFile(os.path.join('data', 'Company Names.xlsx'))
    sheet_names = xl.sheet_names

    selected_sector = tk.StringVar()
    selected_sector.set(sheet_names[0])
    start_date_var = tk.StringVar()
    end_date_var = tk.StringVar()

    def on_ok():
        start = start_date_var.get()
        end = end_date_var.get()
        if not (len(start) == 7 and start[4] == '/' and start[:4].isdigit() and start[5:7].isdigit()):
            messagebox.showwarning("Input Error", "Start date must be in yyyy/mm format.")
            return
        if not (len(end) == 7 and end[4] == '/' and end[:4].isdigit() and end[5:7].isdigit()):
            messagebox.showwarning("Input Error", "End date must be in yyyy/mm format.")
            return
        if int(end[:4] + end[5:7]) < int(start[:4] + start[5:7]):
            messagebox.showwarning("Input Error", "End date must not be before start date.")
            return
        root.selected_sector = selected_sector.get()
        root.start_date = start
        root.end_date = end
        root.destroy()

    label = tk.Label(root, text="Select sector:")
    label.pack(pady=(20, 5))
    dropdown = ttk.Combobox(root, textvariable=selected_sector, values=sheet_names, state="readonly")
    dropdown.pack(pady=5)

    start_label = tk.Label(root, text="Start date (yyyy/mm):")
    start_label.pack(pady=(10, 2))
    start_entry = tk.Entry(root, textvariable=start_date_var)
    start_entry.pack(pady=2)

    end_label = tk.Label(root, text="End date (yyyy/mm):")
    end_label.pack(pady=(10, 2))
    end_entry = tk.Entry(root, textvariable=end_date_var)
    end_entry.pack(pady=2)

    ok_btn = tk.Button(root, text="OK", command=on_ok)
    ok_btn.pack(pady=15)

    root.mainloop()
    return getattr(root, 'selected_sector', None), getattr(root, 'start_date', None), getattr(root, 'end_date', None)

def select_report_options():
    root = tk.Tk()
    root.title("Select Report Options")
    root.geometry("350x220")
    root.resizable(False, False)

    options = [
        ("Z-Score Matrix", "zscore"),
        ("Earnings vs Dididend Plots", "earnings_dividend"),
        ("Relative Graphs", "relative"),
        ("Individual Analysis", "individual"),
    ]
    vars = {}
    for idx, (label, key) in enumerate(options):
        var = tk.BooleanVar()
        chk = tk.Checkbutton(root, text=label, variable=var)
        chk.pack(anchor='w', padx=30, pady=5)
        vars[key] = var

    def on_ok():
        root.selected_options = {k: v.get() for k, v in vars.items()}
        root.destroy()

    ok_btn = tk.Button(root, text="OK", command=on_ok)
    ok_btn.pack(pady=15)

    root.mainloop()
    return getattr(root, 'selected_options', {})

def add_word_toc(doc):
    """Insert a Word TOC field code that will become a clickable TOC when opened in Word."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    doc.add_page_break()

if __name__ == "__main__":
    sector, start_date, end_date = select_sector_and_dates()
    selected_options = select_report_options()
    print("Sector:", sector)
    print("Start Date:", start_date)
    print("End Date:", end_date)
    print("Selected Options:", selected_options)
    print("Generating report...")

    today_str = datetime.datetime.now().strftime("%d%m%Y")
    output_dir = os.path.join(os.path.dirname(__file__), "Reports")
    os.makedirs(output_dir, exist_ok=True)
    doc_path = os.path.join(output_dir, f"{sector.lower()}_{today_str}.docx")
    doc = Document()

    # --- Cover Page (Page 1) ---
    doc.add_heading(f"{sector} Report", 0)
    doc.add_paragraph(f"Date range: {start_date} to {end_date}")
    doc.add_paragraph(f"Report generated on: {today_str[:2]}/{today_str[2:4]}/{today_str[4:]}")
    doc.add_page_break()

    # --- Table of Contents (Page 2, hyperlinked) ---
    doc.add_heading("Table of Contents", level=1)
    add_word_toc(doc)  # This will be a clickable TOC after updating in Word

    numerators = []
    individual_tickers = []

    # --- Main Content ---
    # 1. Comparative Z-Score Matrix
    if selected_options.get("zscore"):
        sector_analysis_dir = os.path.join(os.path.dirname(__file__), "Sector Analysis")
        zscore_path = os.path.join(sector_analysis_dir, "sector_z-scorematrix.py")
        spec = importlib.util.spec_from_file_location("sector_zscorematrix", zscore_path)
        zscore_module = importlib.util.module_from_spec(spec)
        sys.modules["sector_zscorematrix"] = zscore_module
        spec.loader.exec_module(zscore_module)

        heatmap_path = zscore_module.produce_zscore_matrix(sector, start_date, end_date)
        doc.add_heading("1. Comparative Z-Score Matrix", level=1)
        doc.add_picture(heatmap_path, width=Inches(6))
        doc.add_paragraph(f"Date range: {start_date} to {end_date}")

    # 2. Earnings vs Dividend Plots
    if selected_options.get("earnings_dividend"):
        sector_analysis_dir = os.path.join(os.path.dirname(__file__), "Sector Analysis")
        earn_vs_div_path = os.path.join(sector_analysis_dir, "sector_earn_vs_div_plots.py")
        spec = importlib.util.spec_from_file_location("sector_earn_vs_div_plots", earn_vs_div_path)
        earn_vs_div_module = importlib.util.module_from_spec(spec)
        sys.modules["sector_earn_vs_div_plots"] = earn_vs_div_module
        spec.loader.exec_module(earn_vs_div_module)

        plot1_path, plot2_path = earn_vs_div_module.produce_earnings_vs_div_plots(sector, start_date, end_date)
        doc.add_heading("2. Earnings vs Dividend Plots", level=1)
        doc.add_heading('2.1 Z-score P/E vs D/Y', level=2)
        doc.add_picture(plot1_path, width=Inches(6))
        doc.add_heading('2.2 Abs P/E vs Abs D/Y', level=2)
        doc.add_picture(plot2_path, width=Inches(6))

    # 3. Relative Analysis
    if selected_options.get("relative"):
        sector_analysis_dir = os.path.join(os.path.dirname(__file__), "Sector Analysis")
        relative_figures_path = os.path.join(sector_analysis_dir, "sector_relative_figures.py")
        spec = importlib.util.spec_from_file_location("sector_relative_figures", relative_figures_path)
        relative_figures_module = importlib.util.module_from_spec(spec)
        sys.modules["sector_relative_figures"] = relative_figures_module
        spec.loader.exec_module(relative_figures_module)

        plots = relative_figures_module.produce_relative_figures(sector, start_date, end_date)
        grouped = defaultdict(list)
        for pair_name, plot_path in plots:
            numerator = pair_name.split(" / ")[0]
            grouped[numerator].append((pair_name, plot_path))
        numerators = list(grouped.keys())
        doc.add_heading("3. Relative Analysis", level=1)
        for idx, numerator in enumerate(numerators, 1):
            doc.add_heading(f"3.{idx}. {numerator}", level=2)
            for pair_name, plot_path in grouped[numerator]:
                doc.add_heading(f"Relative Analysis: {pair_name}", level=3)
                doc.add_picture(plot_path, width=Inches(6))

    # 4. Individual Analysis
    if selected_options.get("individual"):
        sector_analysis_dir = os.path.join(os.path.dirname(__file__), "Sector Analysis")
        individual_analysis_path = os.path.join(sector_analysis_dir, "sector_individual_analysis.py")
        spec = importlib.util.spec_from_file_location("sector_individual_analysis", individual_analysis_path)
        individual_analysis_module = importlib.util.module_from_spec(spec)
        sys.modules["sector_individual_analysis"] = individual_analysis_module
        spec.loader.exec_module(individual_analysis_module)

        plots = individual_analysis_module.produce_individual_analysis(sector, start_date, end_date)
        individual_tickers = [ticker for ticker, _ in plots]
        doc.add_heading("4. Individual Analysis", level=1)
        for idx, (ticker, plot_path) in enumerate(plots, 1):
            doc.add_heading(f"4.{idx}. {ticker}", level=2)
            doc.add_picture(plot_path, width=Inches(6))

    doc.save(doc_path)
    print(f"Word report saved to {doc_path}")

    # Inform the user with a popup window
    tk.Tk().withdraw()  # Hide the root window
    messagebox.showinfo("Report Generated", "Report generated successfully.")